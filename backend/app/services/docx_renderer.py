import io
import json
from collections import defaultdict
from typing import Any

from docx import Document

_ACRONYMS = {"car", "otd", "ncr", "sla", "llm", "sqr", "id", "duns", "vat"}


def _fmt(id_: str) -> str:
    return " ".join(
        w.upper() if w.lower() in _ACRONYMS else w.capitalize()
        for w in id_.split("_")
    )


def _parse_value(raw: str | None) -> Any:
    if not raw:
        return None
    t = raw.strip()
    if t.startswith("[") or t.startswith("{"):
        try:
            return json.loads(t)
        except json.JSONDecodeError:
            pass
    return raw


def _add_list_table(doc: Document, rows: list[dict]) -> None:
    if not rows:
        doc.add_paragraph("(empty table)")
        return
    keys = [k for k in rows[0] if k != "id"]
    table = doc.add_table(rows=1 + len(rows), cols=len(keys))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, k in enumerate(keys):
        hdr[i].text = _fmt(k)
        run = hdr[i].paragraphs[0].runs
        if run:
            run[0].bold = True
    for ri, row in enumerate(rows, 1):
        cells = table.rows[ri].cells
        for ci, k in enumerate(keys):
            val = row.get(k)
            cells[ci].text = ", ".join(str(v) for v in val) if isinstance(val, list) else str(val if val is not None else "")
    doc.add_paragraph()


def _add_dict_table(doc: Document, obj: dict) -> None:
    entries = [(k, v) for k, v in obj.items() if v is not None and v != ""]
    if not entries:
        return
    table = doc.add_table(rows=len(entries), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(entries):
        cells = table.rows[i].cells
        cells[0].text = _fmt(k)
        run = cells[0].paragraphs[0].runs
        if run:
            run[0].bold = True
        cells[1].text = ", ".join(str(x) for x in v) if isinstance(v, list) else str(v)
    doc.add_paragraph()


def generate_docx(report: dict, fields: list[dict]) -> bytes:
    doc = Document()

    doc.add_heading("Supplier Qualification Report", 0)

    intake = report.get("intake_data") or {}
    supplier = str(intake.get("supplier_name", "—"))
    score = report.get("score")
    score_str = f"{score:.2f} / 5.00" if score is not None else "—"
    verdict = report.get("verdict") or "—"
    evaluator = str(intake.get("evaluator_name", "—"))

    sub = doc.add_paragraph()
    sub.add_run(supplier).bold = True
    sub.add_run(f"  ·  {verdict}  ·  Score {score_str}  ·  {evaluator}")
    doc.add_paragraph()

    sections: dict[str, list[dict]] = defaultdict(list)
    for f in sorted(fields, key=lambda x: x.get("field_index", 0)):
        sections[f["section_id"]].append(f)

    for section_id, sec_fields in sections.items():
        doc.add_heading(_fmt(section_id), level=1)
        for f in sec_fields:
            label = _fmt(f["field_id"])
            parsed = _parse_value(f.get("value"))
            if isinstance(parsed, list) and parsed and isinstance(parsed[0], dict):
                p = doc.add_paragraph()
                p.add_run(label).bold = True
                _add_list_table(doc, parsed)
            elif isinstance(parsed, dict):
                p = doc.add_paragraph()
                p.add_run(label).bold = True
                _add_dict_table(doc, parsed)
            else:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(parsed) if parsed else "—")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
